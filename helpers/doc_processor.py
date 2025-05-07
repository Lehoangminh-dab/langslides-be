import os
import io
import uuid
import tempfile
import logging
from typing import List, Dict, Any, Tuple
from docx import Document
from PIL import Image
from sentence_transformers import SentenceTransformer

logger = logging.getLogger(__name__)

class DocProcessor:
    """
    Process DOC/DOCX files to extract text content and images
    """
    
    def __init__(self, image_embedding_model=None):
        """
        Initialize the DocProcessor with an optional image embedding model
        
        :param image_embedding_model: Model to create embeddings for images (SentenceTransformer)
        """
        self.image_embedding_model = image_embedding_model
        
    def process_doc_file(self, file_content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a DOC/DOCX file and extract text content and images
        
        :param file_content: Raw DOC/DOCX file content as bytes
        :return: Tuple of (extracted_text, extracted_images)
        """
        # Create temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(file_content)
        temp_file.close()
        
        # Create temp directory for extracted images
        img_dir = os.path.join(tempfile.gettempdir(), f"doc_images_{uuid.uuid4().hex}")
        os.makedirs(img_dir, exist_ok=True)
        
        extracted_text = []
        extracted_images = []
        image_count = 0
        
        try:
            # Open the document
            doc = Document(temp_file.name)
            
            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    extracted_text.append(para.text)
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_content = self._extract_table(table)
                if table_content:
                    extracted_text.append(f"\n\n[Table {i+1}]\n{table_content}\n\n")
            
            # Process images
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        # Get the image bytes
                        image_bytes = rel.target_part.blob
                        
                        # Generate a filename (before incrementing count)
                        img_ext = rel.target_ref.split('.')[-1]
                        img_filename = f"image_{image_count}.{img_ext}"
                        img_path = os.path.join(img_dir, img_filename)
                        
                        # Save image to temp file
                        with open(img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        # Process image if embedding model available
                        if self.image_embedding_model:
                            try:
                                # Open with PIL for processing
                                pil_img = Image.open(io.BytesIO(image_bytes))
                                
                                # Log the image dimensions for debugging
                                logger.info(f"Processing image {img_filename}: {pil_img.width}x{pil_img.height}")
                                
                                # Skip very small images (likely icons, etc.)
                                if pil_img.width < 100 or pil_img.height < 100:
                                    logger.info(f"Skipping small image {img_filename}: {pil_img.width}x{pil_img.height}")
                                    continue
                                
                                # Create embedding for the image
                                logger.info(f"Creating embedding for image {img_filename}")
                                img_embedding = self.image_embedding_model.encode(pil_img)
                                
                                # Check if embedding was created successfully
                                if img_embedding is None or len(img_embedding) == 0:
                                    logger.warning(f"Failed to create embedding for image {img_filename}")
                                    continue
                                
                                # Store image with metadata
                                extracted_images.append({
                                    "path": img_path,
                                    "image_bytes": image_bytes,
                                    "embedding": img_embedding,
                                    "width": pil_img.width,
                                    "height": pil_img.height,
                                    "aspect_ratio": pil_img.width / pil_img.height
                                })
                                
                                # Add image reference to text
                                extracted_text.append(f"\n[Image {image_count}: {img_filename}]\n")
                                logger.info(f"Successfully processed image {img_filename}")
                                
                            except Exception as e:
                                logger.warning(f"Failed to process image {img_filename}: {e}", exc_info=True)
                        
                        # Always increment the counter even if processing failed
                        image_count += 1
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}", exc_info=True)
            
            # Combine all extracted text
            full_text = "\n".join(extracted_text)
            logger.info(f"Extracted {len(extracted_text)} text blocks and {len(extracted_images)} images")
            
            return full_text, extracted_images
            
        except Exception as e:
            logger.error(f"Error processing DOC file: {e}", exc_info=True)
            return "", []
        finally:
            # Clean up the temp file
            os.unlink(temp_file.name)
    
    def _extract_table(self, table) -> str:
        """
        Extract table content as markdown, avoiding namespace issues
        
        :param table: docx Table object
        :return: Markdown representation of the table
        """
        table_text = []
        
        # Process each row
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell_idx, cell in enumerate(row.cells):
                # Get cell text without checking for images (to avoid namespace issues)
                cell_text = self._get_cell_text(cell)
                row_text.append(cell_text)
            
            table_text.append(" | ".join(row_text))
        
        # Format as markdown table
        if table_text:
            if len(table_text) > 1:  # Has at least a header and a data row
                header_row = table_text[0]
                col_count = len(header_row.split("|"))
                separator_row = " | ".join(["-" * 3 for _ in range(col_count)])
                markdown_table = [
                    header_row,
                    separator_row
                ] + table_text[1:]
                
                return "\n".join(markdown_table)
            else:  # Only one row
                return table_text[0]
        
        return ""
    
    def _get_cell_text(self, cell) -> str:
        """
        Extract text from a table cell, including handling nested paragraphs
        """
        return "\n".join([p.text for p in cell.paragraphs if p.text.strip()])